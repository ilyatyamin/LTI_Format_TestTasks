import json
import zipfile
import xml.etree.ElementTree as ElTree
from docx import Document

from scripts.fromLTI.PlatformsMultipleChoice import PlatformsMultipleChoice
from scripts.toLTI.MultipleChoiceQuestion import MultiplyChoiceQuestion
from scripts.toLTI.Question import Question
from scripts.toLTI.conversion_formats import ConversionFormat
from scripts.web_requests.RequestManager import RequestManager
from scripts.web_requests.RequestPlatform import RequestPlatform


class FormatsHandler:
    """
    The base class for handling all transformations in system
    """
    def __init__(self):
        self.__manager_multiple_choice = MultipleChoiceManager()
        self.__platforms_multiple_choice = PlatformsMultipleChoice()
        self.xml_formats = [int(ConversionFormat.MultipleChoiceMoodleXML)]
        self.json_formats = [int(ConversionFormat.LTI), int(ConversionFormat.MultipleChoiceStepikStep),
                             int(ConversionFormat.CanvasInstructure)]
        self.csv_formats = [int(ConversionFormat.MultipleChoiceMoodleCSV)]
        self.docx_formats = [int(ConversionFormat.MultipleChoiceTestmozWord),
                             int(ConversionFormat.MultipleChoiceMoodleWord)]

    def get_format(self, current_format: int):
        if current_format in self.xml_formats:
            return 'xml'
        if current_format in self.json_formats:
            return 'json'
        if current_format in self.csv_formats:
            return 'csv'
        if current_format in self.docx_formats:
            return 'docx'

    def process_file_based_question(self, path_to_file: str, file_structure: int, needed_format: int):
        """
        Method get path_to_file in string representations, it's file structure (Conversion Format enum) and format in which user need's to convert his/her file.
        Returns MultipleChoice / list[MultipleChoice] / str in dependent of types
        """
        if ConversionFormat.is_multiple_choice(int(file_structure)):
            question = self.__manager_multiple_choice.process_question_lti(path_to_file, file_structure)
            answer = self.__tranform_from_lti_to_platform(question, needed_format)
            return answer

    def process_request_based_question(self, courseId: int, quizId: int, file_structure: int, needed_format: int):
        if ConversionFormat.is_multiple_choice(int(file_structure)):
            if file_structure == int(ConversionFormat.CanvasInstructure):
                question = self.__manager_multiple_choice.process_question_lti('', file_structure, courseId, quizId)
                return self.__tranform_from_lti_to_platform(question, needed_format)

    def __tranform_from_lti_to_platform(self, question: Question, format_creat: int):
        """
        Transforms LTI created question to platform's format.
        If format_creat is LTI, method returns string-presentation of question in LTI-style
        """
        if format_creat == int(ConversionFormat.LTI):
            return self.get_text(question)
        else:
            return self.__platforms_multiple_choice.parse_questions(question, format_creat)

    def get_text(self, obj) -> str:
        """
        Return string-presentation of question.
        Nevertheless, this method works if you will put in him string: it will return it without any transformations
        """
        if isinstance(obj, Question) or isinstance(obj, list):
            return json.dumps(obj, default=lambda o: self.encoder(o.__dict__), ensure_ascii=False, indent=4)
        if isinstance(obj, str):
            return obj

    @staticmethod
    def encoder(dictionary: dict):
        """
        Smart-encoding method. It doesn't take empty attributes in classes, just filled.
        """
        values_to_delete = []
        for item, value in dictionary.items():
            if (value is None or (isinstance(value, str) and (value.isspace()))
                    or (isinstance(value, list) and len(value) == 0)
                    or (isinstance(value, list) and len(value) != 0 and isinstance(value[0], str) and all(
                        x.isspace() for x in value))):
                values_to_delete.append(item)
        for value in values_to_delete:
            dictionary.pop(value)
        return dictionary

    def write_to_file(self, name_of_file: str, obj):
        """
        Write text-presentation of question to file
        """
        f = open(name_of_file, 'w+')
        f.seek(0)

        if isinstance(obj, MultiplyChoiceQuestion) or isinstance(obj, list):
            json.dump(obj, f, default=lambda o: self.encoder(o.__dict__), ensure_ascii=False, indent=4)
        elif isinstance(obj, str):
            f.write(obj)
        f.truncate()
        f.close()

    def write_to_archive(self, name_of_file: str, obj: list[str], format_writing: str):
        """
        Creates an archive in the name_of_file path.
        In this archive writes files in questions, one question for one file.
        """
        i = 0
        with zipfile.ZipFile(name_of_file, 'w') as zipf:
            # Добавляем каждый файл в архив
            for file in obj:
                zipf.writestr(f'question_{i}.{format_writing}', file)
                i += 1



class Manager:

    def open_file(self, path: str) -> list[str]:
        with open(path, encoding='utf8') as file:
            return file.readlines()

    def strip_file(self, path_to_file: str):
        lines = []
        with open(path_to_file, 'r+') as file:
            for line in file:
                if not line.isspace():
                    lines.append(line)
        with open(path_to_file, 'r+') as file:
            for line in lines:
                file.write(line)
            file.truncate()
            file.close()


class MultipleChoiceManager(Manager):
    """
    Load banks of files
    """

    def process_question_lti(self, path_to_file: str, file_structure: int, courseId=None, quizId=None):
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleXML):
            # check one or bank of questions
            self.strip_file(path_to_file)
            parsed_file = [x for x in ElTree.parse(path_to_file).getroot().findall("question") if
                           x.attrib['type'] == 'multichoice']
            if len(parsed_file) > 0:
                return self.__load_bank_question__(path_to_file, file_structure)
            else:
                return self.__load_one_question__(path_to_file, file_structure)
        elif file_structure == int(ConversionFormat.MultipleChoiceMoodleWord):
            doc = Document(path_to_file)
            num = len(doc.tables)
            if num == 1:
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_moodle_docx(doc, 0)
                return obj
            else:
                bank = []
                for i in range(num):
                    obj = MultiplyChoiceQuestion()
                    obj.parse_one_question_moodle_docx(doc, i)
                    bank.append(obj)
                return bank
        elif file_structure == int(ConversionFormat.MultipleChoiceStepikStep):
            # T-O-D-O: как парсить архив?
            return self.__load_one_question__(path_to_file, file_structure)
        elif file_structure == int(ConversionFormat.MultipleChoiceTestmozWord):
            # T-O-D-O: как парсить много вопросов?
            return self.__load_one_question__(path_to_file, file_structure)
        elif file_structure == int(ConversionFormat.CanvasInstructure):
            request = RequestManager.get_quiz_info(courseId, quizId, RequestPlatform.Canvas)
            if request.status_code != 200:
                raise Exception(
                    'You have no rights for this. Please sure that you add user as teacher in your course ' +
                    'and you correctly type courseId and quizId.')
            bank = []
            for question in json.loads(request.text):
                if question['question_type'] == 'multiple_choice_question':
                    obj = MultiplyChoiceQuestion()
                    obj.parse_one_question_canvas(question)
                    bank.append(obj)
            return bank

    def __load_one_question__(self, path_to_file: str, file_structure: int) -> MultiplyChoiceQuestion:
        obj = MultiplyChoiceQuestion()
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleXML):
            self.strip_file(path_to_file)
            parsed_file = ElTree.parse(path_to_file).getroot()
            file_as_txt = open(path_to_file).readlines()
            obj.parse_one_question_from_moodle_xml(parsed_file)
            return obj
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleCSV):
            lines = [x.split(',') for x in self.open_file(path_to_file) if not (x.isspace())]
            idx = 1
            obj.parse_one_question_moodle_csv(lines, idx)
            return obj
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleWord):
            doc = Document(path_to_file)
            obj.parse_one_question_moodle_docx(doc, 0)
            return obj
        if file_structure == int(ConversionFormat.MultipleChoiceStepikStep):
            with open(path_to_file) as file:
                parsed_file = json.load(file)
                obj.parse_one_question_stepik(parsed_file)
            return obj
        if file_structure == int(ConversionFormat.MultipleChoiceTestmozWord):
            doc = Document(path_to_file)
            obj.parse_one_question_testmoz_word(doc, 0)
            return obj
        return obj

    def __load_bank_question__(self, path_to_file: str, file_structure: int):
        bank = []
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleWord):
            doc = Document(path_to_file)
            num = len(doc.tables)
            for i in range(num):
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_moodle_docx(doc, i)
                bank.append(obj)
            return bank
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleXML):
            self.strip_file(path_to_file)
            parsed_file = [x for x in ElTree.parse(path_to_file).getroot().findall("question") if
                           x.attrib['type'] == 'multichoice']
            for question in parsed_file:
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_from_moodle_xml(question)
                bank.append(obj)
            return bank
        if file_structure == int(ConversionFormat.MultipleChoiceMoodleCSV):
            lines = [x.split(',') for x in self.open_file(path_to_file) if not (x.isspace())]
            for idx in range(1, len(lines)):
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_moodle_csv(lines, idx)
                bank.append(obj)
            return bank
        if file_structure == int(ConversionFormat.MultipleChoiceStepikStep):
            """
            TODO THIS! 
            подумать как сделать импорт набора stepik вопросов: наверное, все же, архив?
            """
        return bank
