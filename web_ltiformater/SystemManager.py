import enum
import json
import xml.etree.ElementTree as ElTree
from docx import Document


class SystemManager:
    class MultipleChoiceEnum(enum.Enum):
        MoodleXml = 0  # ready
        MoodleWord = 1  # ready
        MoodleAiken = 2
        MoodleCSV = 3  # ready
        MoodleGift = 4
        SocrativePDF = 5
        StepikStep = 6  # ready
        YaContestJSON = 7
        OnlineTestPadPDF = 8
        TestMozWord = 9  # ready
        TestMozCSV = 10
        ImportSocrativeExcel = 11
        ImportQuizizzExcel = 12
        ImportKahootExcel = 13
        ImportBlooketExcel = 14
        ImportTestMozExcel = 15
        ImportProProofsExcel = 16
        ImportClassMakerExcel = 17

    def load_one_question(self, path_to_file: str, file_structure: enum.Enum):
        pass

    def load_bank_question(self, path_to_file: str, file_structure: enum.Enum):
        pass

    def write_to_file(self, name_of_file: str, obj):
        pass

    def encoder(self, dictionary: dict):
        values_to_delete = []
        for item, value in dictionary.items():
            if (value is None or (isinstance(value, str) and (value.isspace()))
                    or (isinstance(value, list) and len(value) == 0)
                    or (isinstance(value, list) and len(value) != 0 and isinstance(value[0], str) and all(
                        x.isspace() for x in value))):
                values_to_delete.append(item)
        for value in values_to_delete:
            dictionary.pop(value)
        return dictionary

    def strip_file(self, path_to_file: str):
        lines = []
        with open(path_to_file, 'r+') as file:
            for line in file:
                if not line.isspace():
                    lines.append(line)
        with open(path_to_file, 'r+') as file:
            for line in lines:
                file.write(line)
            file.truncate()
            file.close()

    def open_file(self, path: str) -> list[str]:
        with open(path, encoding='utf8') as file:
            return file.readlines()


class MultipleChoiceManager(SystemManager):
    """
    Load banks of files
    """

    def load_one_question(self, path_to_file: str, file_structure: enum.Enum) -> MultiplyChoiceQuestion:
        obj = MultiplyChoiceQuestion()
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleXml:
            self.strip_file(path_to_file)
            parsed_file = ElTree.parse(path_to_file).getroot()
            file_as_txt = open(path_to_file).readlines()
            obj.parse_one_question_from_moodle_xml(parsed_file)
            return obj
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleCSV:
            lines = [x.split(',') for x in self.open_file(path_to_file) if not (x.isspace())]
            idx = 1
            obj.parse_one_question_moodle_csv(lines, idx)
            return obj
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleWord:
            doc = Document(path_to_file)
            obj.parse_one_question_moodle_docx(doc, 0)
            return obj
        if file_structure == SystemManager.MultipleChoiceEnum.StepikStep:
            with open(path_to_file) as file:
                parsed_file = json.load(file)
                obj.parse_one_question_stepik(parsed_file)
            return obj
        if file_structure == SystemManager.MultipleChoiceEnum.TestMozWord:
            doc = Document(path_to_file)
            obj.parse_one_question_testmoz_word(doc, 3)
            return obj
        return obj

    def load_bank_question(self, path_to_file: str, file_structure: enum.Enum):
        bank = []
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleWord:
            doc = Document(path_to_file)
            num = len(doc.tables)
            for i in range(num):
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_moodle_docx(doc, i)
                bank.append(obj)
            return bank
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleXml:
            self.strip_file(path_to_file)
            parsed_file = [x for x in ElTree.parse(path_to_file).getroot().findall("question") if
                           x.attrib['type'] == 'multichoice']
            for question in parsed_file:
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_from_moodle_xml(question)
                bank.append(obj)
            return bank
        if file_structure == SystemManager.MultipleChoiceEnum.MoodleCSV:
            lines = [x.split(',') for x in self.open_file(path_to_file) if not (x.isspace())]
            for idx in range(1, len(lines)):
                obj = MultiplyChoiceQuestion()
                obj.parse_one_question_moodle_csv(lines, idx)
                bank.append(obj)
            return bank
        if file_structure == SystemManager.MultipleChoiceEnum.StepikStep:
            """
            TODO THIS! 
            подумать как сделать импорт набора stepik вопросов: наверное, все же, архив?
            """
        return bank

    def write_to_file(self, name_of_file: str, obj):
        f = open(name_of_file, 'w+')
        f.seek(0)

        # Serialize the data and write it to the file
        json.dump(obj, f, default=lambda o: self.encoder(o.__dict__), ensure_ascii=False, indent=4)
        f.truncate()
        f.close()
